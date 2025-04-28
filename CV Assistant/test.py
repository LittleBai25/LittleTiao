import streamlit as st
import os
import json
import requests
import pandas as pd
import io
import base64
from tempfile import NamedTemporaryFile
from PIL import Image

# 设置页面标题和配置
st.set_page_config(
    page_title="个人简历写作助理",
    page_icon="📝",
    layout="wide"
)

# 设置页面样式
st.markdown("""
<style>
    .main {
        padding: 2rem;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 2rem;
    }
    .stTabs [data-baseweb="tab"] {
        font-size: 1rem;
        font-weight: 600;
    }
    .upload-box {
        border: 2px dashed #4e8df5;
        border-radius: 10px;
        padding: 1.5rem;
        margin-bottom: 1rem;
    }
    .model-selector {
        margin-bottom: 2rem;
    }
    .prompt-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 5px;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# 从Streamlit Secrets获取API密钥
def get_api_key():
    try:
        return st.secrets["openrouter_api_key"]
    except Exception:
        # 本地开发时可能没有设置secrets
        return None

# 读取Excel文件
def read_excel(uploaded_file):
    if uploaded_file is not None:
        try:
            return pd.read_excel(uploaded_file)
        except Exception as e:
            st.error(f"无法读取Excel文件: {e}")
            return None
    return None

# 读取文本文件
def read_text_file(uploaded_file):
    if uploaded_file is not None:
        try:
            return uploaded_file.getvalue().decode("utf-8")
        except Exception as e:
            st.error(f"无法读取文本文件: {e}")
            return None
    return None

# 读取PDF文件
def read_pdf(uploaded_file):
    if uploaded_file is not None:
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.getvalue()))
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"无法读取PDF文件: {e}")
            return None
    return None

# 增强版Word文件读取，改进对表格的支持
def read_docx(uploaded_file):
    if uploaded_file is not None:
        try:
            import docx
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            
            # 提取段落内容
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():  # 忽略空段落
                    text += para.text + "\n\n"
            
            # 增强的表格提取
            for i, table in enumerate(doc.tables):
                text += f"\n### 表格 {i+1}:\n"
                
                # 计算每一列的最大宽度，以便对齐
                col_widths = []
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if i >= len(col_widths):
                            col_widths.append(len(cell_text))
                        else:
                            col_widths[i] = max(col_widths[i], len(cell_text))
                
                # 创建表格的Markdown格式
                # 添加表头行
                first_row = table.rows[0]
                header = "| "
                for i, cell in enumerate(first_row.cells):
                    cell_text = cell.text.strip()
                    padding = col_widths[i] - len(cell_text)
                    header += cell_text + " " * padding + " | "
                text += header + "\n"
                
                # 添加分隔行
                separator = "| "
                for width in col_widths:
                    separator += "-" * width + " | "
                text += separator + "\n"
                
                # 添加内容行（从第二行开始）
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:  # 跳过表头
                        continue
                    row_text = "| "
                    for i, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        padding = col_widths[i] - len(cell_text)
                        row_text += cell_text + " " * padding + " | "
                    text += row_text + "\n"
                
                text += "\n"  # 表格后添加空行
            
            return text
        except Exception as e:
            st.error(f"无法读取Word文件: {e}")
            return None
    return None

# 读取图片文件
def read_image(uploaded_file):
    if uploaded_file is not None:
        try:
            # 打开图片
            image = Image.open(io.BytesIO(uploaded_file.getvalue()))
            
            # 转换为base64以便在应用中显示
            buffered = io.BytesIO()
            image_format = image.format if image.format else "PNG"
            image.save(buffered, format=image_format)
            img_str = base64.b64encode(buffered.getvalue()).decode()
            
            # 返回图片信息
            return {
                "type": "image",
                "format": image_format,
                "size": image.size,
                "mode": image.mode,
                "base64": img_str
            }
        except Exception as e:
            st.error(f"无法读取图片文件: {e}")
            return None
    return None

# 处理上传的文件
def process_file(uploaded_file):
    if uploaded_file is None:
        return None
    
    file_extension = uploaded_file.name.split(".")[-1].lower()
    
    # 根据文件类型处理
    if file_extension in ["xlsx", "xls"]:
        return read_excel(uploaded_file)
    elif file_extension == "pdf":
        return read_pdf(uploaded_file)
    elif file_extension in ["docx", "doc"]:
        return read_docx(uploaded_file)
    elif file_extension in ["txt", "md"]:
        return read_text_file(uploaded_file)
    elif file_extension in ["jpg", "jpeg", "png", "gif", "bmp"]:
        return read_image(uploaded_file)
    else:
        st.warning(f"不支持的文件类型: .{file_extension}")
        return None

# 调用OpenRouter API
def call_openrouter_api(model, messages):
    api_key = get_api_key()
    
    if not api_key:
        st.error("未找到API密钥。请确保已在Streamlit的secrets中设置了openrouter_api_key。")
        return None
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://resume-assistant.streamlit.app",  # 你的应用URL
        "X-Title": "个人简历写作助理"  # 你的应用名称
    }
    
    data = {
        "model": model,
        "messages": messages
    }
    
    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            data=json.dumps(data)
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            st.error(f"API调用失败: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        st.error(f"API调用出错: {e}")
        return None

# 主应用
def main():
    st.title("个人简历写作助理")
    
    # 初始化会话状态
    if 'resume_data' not in st.session_state:
        st.session_state['resume_data'] = None
    if 'resume_file_name' not in st.session_state:
        st.session_state['resume_file_name'] = None
    if 'support_data' not in st.session_state:
        st.session_state['support_data'] = {}
    if 'selected_model' not in st.session_state:
        st.session_state['selected_model'] = "anthropic/claude-3-5-sonnet"
    if 'persona' not in st.session_state:
        st.session_state['persona'] = ""
    if 'task' not in st.session_state:
        st.session_state['task'] = ""
    if 'output_format' not in st.session_state:
        st.session_state['output_format'] = ""
    
    # 创建两个标签页
    tab1, tab2 = st.tabs(["📄 上传简历素材", "⚙️ 模型设置与提示词"])
    
    # 第一个标签页：文件上传
    with tab1:
        st.header("上传简历素材")
        
        # 个人简历素材表上传（单选，必传）
        st.subheader("个人简历素材表（单选，必传）")
        st.markdown("请上传包含您个人信息的文件（Excel、Word、PDF等格式均可）。")
        
        with st.container():
            st.markdown('<div class="upload-box">', unsafe_allow_html=True)
            resume_file = st.file_uploader("选择您的个人简历素材表", 
                                          type=["xlsx", "xls", "pdf", "docx", "doc", "txt", "md"], 
                                          key="resume_file")
            
            if resume_file is None:
                st.warning("⚠️ 个人简历素材表是必须上传的")
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            if resume_file is not None:
                # 处理简历文件
                resume_data = process_file(resume_file)
                if resume_data is not None:
                    st.success(f"成功上传: {resume_file.name}")
                    
                    # 根据数据类型显示不同的预览
                    with st.expander("预览简历素材表"):
                        if isinstance(resume_data, pd.DataFrame):
                            st.dataframe(resume_data)
                        elif isinstance(resume_data, dict) and "type" in resume_data and resume_data["type"] == "image":
                            st.image(f"data:image/{resume_data['format'].lower()};base64,{resume_data['base64']}")
                            st.text(f"图片信息: 格式 {resume_data['format']}, 尺寸 {resume_data['size'][0]}x{resume_data['size'][1]}")
                        else:
                            # 纯文本
                            st.markdown(resume_data[:5000] + "..." if len(resume_data) > 5000 else resume_data)
                    
                    # 将数据保存到会话状态
                    st.session_state['resume_data'] = resume_data
                    st.session_state['resume_file_name'] = resume_file.name
        
        # 支持文件上传（多选，非必传）
        st.subheader("支持文件（多选，非必传）")
        st.markdown("上传相关的支持文件，如项目描述、工作职责详情、作品集等。")
        
        with st.container():
            st.markdown('<div class="upload-box">', unsafe_allow_html=True)
            support_files = st.file_uploader("选择支持文件", 
                                           type=["pdf", "docx", "doc", "txt", "md", "xlsx", "xls", "jpg", "jpeg", "png", "gif"],
                                           accept_multiple_files=True,
                                           key="support_files")
            st.markdown('</div>', unsafe_allow_html=True)
            
            if support_files:
                st.success(f"成功上传 {len(support_files)} 个支持文件")
                
                support_data = {}
                for file in support_files:
                    file_data = process_file(file)
                    if file_data is not None:
                        support_data[file.name] = file_data
                        with st.expander(f"预览: {file.name}"):
                            if isinstance(file_data, pd.DataFrame):
                                st.dataframe(file_data)
                            elif isinstance(file_data, dict) and "type" in file_data and file_data["type"] == "image":
                                st.image(f"data:image/{file_data['format'].lower()};base64,{file_data['base64']}")
                                st.text(f"图片信息: 格式 {file_data['format']}, 尺寸 {file_data['size'][0]}x{file_data['size'][1]}")
                            else:
                                # 纯文本或Markdown
                                st.markdown(file_data[:3000] + "..." if len(str(file_data)) > 3000 else file_data)
                
                # 将支持文件数据保存到会话状态
                st.session_state['support_data'] = support_data
    
    # 第二个标签页：模型设置和提示词
    with tab2:
        st.header("模型设置与提示词")
        
        # 模型选择
        st.subheader("选择大模型")
        
        with st.container():
            st.markdown('<div class="model-selector">', unsafe_allow_html=True)
            
            # 模型列表
            models = [
                "anthropic/claude-3-5-sonnet",
                "anthropic/claude-3-opus",
                "anthropic/claude-3-haiku",
                "openai/gpt-4-turbo",
                "openai/gpt-4o",
                "openai/gpt-3.5-turbo",
                "google/gemini-pro",
                "google/gemini-1.5-pro",
                "mistralai/mistral-large",
                "mistralai/mistral-7b"
            ]
            
            selected_model = st.selectbox("选择要使用的模型", models, index=models.index(st.session_state['selected_model']) if st.session_state['selected_model'] in models else 0)
            st.session_state['selected_model'] = selected_model
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        # 提示词设置
        st.subheader("提示词设置")
        
        # 默认值
        default_persona = """你是一位专业的简历顾问，擅长根据用户的经历和技能，编写出专业、吸引人的简历内容。你的建议既要突出用户的优势，又要符合行业标准和招聘者的期望。"""
        
        default_task = """请根据提供的个人信息和支持材料，为用户编写一份针对特定职位的简历内容。需要：
1. 分析用户的经历，找出与目标职位最相关的技能和成就
2. 使用具体数据和成果来量化成就
3. 使用行业相关的关键词，以提高通过ATS系统的几率
4. 保持简洁专业的语言风格
5. 根据用户的职业阶段调整内容深度"""
        
        default_format = """输出格式应包含以下部分：
1. 个人信息：姓名、联系方式等基本信息
2. 个人简介：3-5句话概括核心优势和职业目标
3. 工作经历：按照时间倒序排列，包含公司名称、职位、时间段、主要职责和成就
4. 教育背景：学校、专业、学位、毕业时间
5. 技能列表：按照相关性排列的专业技能
6. 项目经验（如适用）：项目名称、角色、时间、描述和成果

每个部分要突出重点，使用简洁有力的语言，总体控制在一页纸内。"""
        
        # 创建三个文本框
        col1, col2 = st.columns(2)
        
        with col1:
            persona = st.text_area("人物设定", 
                                value=st.session_state['persona'] if st.session_state['persona'] else default_persona,
                                height=150)
            st.session_state['persona'] = persona
            
            task = st.text_area("任务描述", 
                              value=st.session_state['task'] if st.session_state['task'] else default_task,
                              height=250)
            st.session_state['task'] = task
        
        with col2:
            output_format = st.text_area("输出格式", 
                                       value=st.session_state['output_format'] if st.session_state['output_format'] else default_format,
                                       height=430)
            st.session_state['output_format'] = output_format
        
        # 生成简历按钮
        st.subheader("生成简历")
        
        if st.button("开始生成简历", type="primary", use_container_width=True):
            # 检查是否上传了简历素材（必传项）
            if 'resume_data' not in st.session_state or st.session_state['resume_data'] is None:
                st.error("请先上传个人简历素材表！这是必须的。")
            else:
                with st.spinner("正在生成您的简历，请稍候..."):
                    # 准备API调用所需的数据
                    resume_data = st.session_state.get('resume_data')
                    resume_file_name = st.session_state.get('resume_file_name', '个人简历素材')
                    support_data = st.session_state.get('support_data', {})
                    
                    # 处理简历数据
                    resume_info = f"个人简历素材表（{resume_file_name}）内容：\n"
                    
                    if isinstance(resume_data, pd.DataFrame):
                        # 如果是DataFrame（Excel）
                        for col in resume_data.columns:
                            resume_info += f"{col}:\n"
                            for idx, value in resume_data[col].items():
                                if pd.notna(value):
                                    resume_info += f"- {value}\n"
                            resume_info += "\n"
                    elif isinstance(resume_data, dict):
                        if "type" in resume_data and resume_data["type"] == "image":
                            # 如果是图片
                            resume_info += f"[这是一张图片文件，格式为{resume_data['format']}，尺寸为{resume_data['size'][0]}x{resume_data['size'][1]}]\n"
                        else:
                            # 其他字典格式
                            for key, value in resume_data.items():
                                if key != "base64":  # 避免输出大量base64数据
                                    resume_info += f"{key}: {value}\n"
                    else:
                        # 普通文本
                        resume_info += str(resume_data)
                    
                    # 整合支持文件内容
                    support_info = ""
                    if support_data:
                        support_info = "支持文件内容：\n"
                        for filename, content in support_data.items():
                            support_info += f"\n--- {filename} ---\n"
                            
                            if isinstance(content, pd.DataFrame):
                                # DataFrame（Excel）
                                support_info += content.to_string()
                            elif isinstance(content, dict):
                                if "type" in content and content["type"] == "image":
                                    # 图片
                                    support_info += f"[这是一张图片文件，格式为{content['format']}，尺寸为{content['size'][0]}x{content['size'][1]}]\n"
                                else:
                                    # 其他字典格式
                                    for key, value in content.items():
                                        if key != "base64":  # 避免输出大量base64数据
                                            support_info += f"{key}: {value}\n"
                            else:
                                # 普通文本
                                support_info += str(content)
                            
                            support_info += "\n"
                    
                    # 构建完整的提示词
                    system_message = f"{persona}\n\n{task}\n\n{output_format}"
                    
                    user_message = f"""请根据以下提供的信息，按照要求编写一份专业的简历：

{resume_info}

{support_info if support_info else '未提供支持文件。'}

请根据以上信息，编写一份专业、有针对性的简历。"""
                    
                    # 准备API调用
                    messages = [
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message}
                    ]
                    
                    # 调用API
                    response = call_openrouter_api(st.session_state['selected_model'], messages)
                    
                    if response:
                        try:
                            result = response['choices'][0]['message']['content']
                            st.session_state['resume_result'] = result
                            
                            # 显示结果
                            st.success("简历生成完成！")
                            st.subheader("生成的简历内容")
                            st.markdown(result)
                            
                            # 提供下载按钮
                            resume_txt = result
                            
                            # 创建临时文件
                            with NamedTemporaryFile(delete=False, suffix='.md') as tmp:
                                tmp.write(resume_txt.encode('utf-8'))
                                tmp_path = tmp.name
                            
                            with open(tmp_path, 'rb') as f:
                                st.download_button(
                                    label="下载简历内容（Markdown）",
                                    data=f,
                                    file_name="简历.md",
                                    mime="text/markdown"
                                )
                            
                            # 删除临时文件
                            os.unlink(tmp_path)
                            
                        except Exception as e:
                            st.error(f"处理API响应时出错: {e}")
                    else:
                        st.error("简历生成失败，请重试。")

if __name__ == "__main__":
    main()